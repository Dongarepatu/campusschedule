# timetable/views.py
import io
import csv
import json
from datetime import datetime
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Department, Faculty, TimetableEntry,TimetableHistory
from .forms import TimetableForm
from django.core.exceptions import ValidationError
from django.http import HttpResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle, SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from django.core import serializers
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
import json
from django.contrib.auth.decorators import login_required


def dashboard(request):
    departments = Department.objects.all()
    return render(request, 'timetable/dashboard.html', {
        'departments': departments,
        'dept_count': departments.count(),
        'faculty_count': Faculty.objects.count(),
    })

def department_list(request):
    departments = Department.objects.all()
    return render(request, 'timetable/department_list.html', {
        'departments': departments
    })

def department_semesters(request, dept_id):
    """Show all semesters available for a department"""
    department = get_object_or_404(Department, id=dept_id)
    
    # Always show all possible semesters based on department type
    if department.name in ['BCA', 'BCS', 'BCOM']:
        semesters = ['Semester 1', 'Semester 2', 'Semester 3', 
                    'Semester 4', 'Semester 5', 'Semester 6']
    elif department.name in ['MCA', 'MCS']:
        semesters = ['Semester 1', 'Semester 2', 'Semester 3', 'Semester 4']
    else:
        semesters = ['Semester 1', 'Semester 2', 'Semester 3', 
                    'Semester 4', 'Semester 5', 'Semester 6']
    
    return render(request, 'timetable/department_semesters.html', {
        'department': department,
        'semesters': semesters,
    })

def timetable_view(request, dept_id):
    """Show timetable for a specific department and semester"""
    department = get_object_or_404(Department, id=dept_id)
    
    # Get selected semester from URL
    selected_semester = request.GET.get('semester', 'Semester 1')
    
    # Get entries for this department and semester
    entries = TimetableEntry.objects.filter(
        department=department,
        semester=selected_semester
    )
    
    # Get all available semesters for this department (always show all)
    if department.name in ['BCA', 'BCS', 'BCOM']:
        all_semesters = ['Semester 1', 'Semester 2', 'Semester 3', 
                        'Semester 4', 'Semester 5', 'Semester 6']
    elif department.name in ['MCA', 'MCS']:
        all_semesters = ['Semester 1', 'Semester 2', 'Semester 3', 'Semester 4']
    else:
        all_semesters = ['Semester 1', 'Semester 2', 'Semester 3', 
                        'Semester 4', 'Semester 5', 'Semester 6']
    
    # Days of the week
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    
    # Get all unique time slots
    if entries.exists():
        unique_times = entries.values_list(
            'start_time', 'end_time'
        ).distinct().order_by('start_time')
    else:
        unique_times = []
    
    # Build timetable matrix with lab session grouping
    matrix = []
    for start_time, end_time in unique_times:
        time_slot = f"{start_time.strftime('%I:%M %p')} - {end_time.strftime('%I:%M %p')}"
        day_data = []
        
        for day in days:
            day_entries = entries.filter(
                day=day,
                start_time=start_time,
                end_time=end_time
            ).select_related('faculty')
            
            # Group entries by subject to identify lab sessions
            if day_entries.count() > 1:
                # Check if all entries have the same subject (lab session)
                subjects = list(day_entries.values_list('subject', flat=True).distinct())
                if len(subjects) == 1:
                    # It's a lab session with multiple teachers
                    lab_subject = subjects[0]
                    lab_faculties = []
                    
                    # Group all lab entries together
                    grouped_entries = []
                    for entry in day_entries:
                        if entry.faculty:
                            lab_faculties.append(entry.faculty.name)
                    grouped_entries.append({
                        'is_lab': True,
                        'subject': lab_subject,
                        'faculties': lab_faculties,
                        'faculty': None,  # No single faculty
                        'id': day_entries.first().id,  # Use first entry ID for deletion
                        'all_ids': list(day_entries.values_list('id', flat=True))  # All IDs for deletion
                    })
                    day_data.append(grouped_entries)
                else:
                    # Different subjects at same time (shouldn't happen but handle it)
                    day_data.append(list(day_entries))
            else:
                # Single entry or empty
                day_data.append(list(day_entries))
        
        matrix.append({
            'time': time_slot,
            'data': day_data
        })
    
    return render(request, 'timetable/timetable_view.html', {
        'department': department,
        'selected_semester': selected_semester,
        'all_semesters': all_semesters,
        'days': days,
        'matrix': matrix,
        'has_entries': entries.exists(),
    })

def download_timetable_pdf(request, dept_id):
    """Generate PDF timetable with proper lab session formatting"""
    department = get_object_or_404(Department, id=dept_id)
    selected_semester = request.GET.get('semester', 'Semester 1')
    
    # Get entries
    entries = TimetableEntry.objects.filter(
        department=department,
        semester=selected_semester
    ).select_related('faculty')
    
    # Create a file-like buffer
    buffer = io.BytesIO()
    
    # Create PDF
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4),
                          rightMargin=20, leftMargin=20,
                          topMargin=40, bottomMargin=40)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#000000'),
        alignment=TA_CENTER,
        spaceAfter=6
    )
    
    dept_style = ParagraphStyle(
        'DeptStyle',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#7c3aed'),
        alignment=TA_CENTER,
        spaceAfter=4
    )
    
    semester_style = ParagraphStyle(
        'SemesterStyle',
        parent=styles['Heading3'],
        fontSize=12,
        textColor=colors.HexColor('#10b981'),
        alignment=TA_CENTER,
        spaceAfter=15
    )
    
    # Title
    elements.append(Paragraph("D.H.B. SONI COLLEGE, SOLAPUR", title_style))
    elements.append(Paragraph(f"{department.name} Department", dept_style))
    elements.append(Paragraph(f"{selected_semester} Timetable", semester_style))
    
    # Days
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    
    # Get unique time slots
    if entries.exists():
        unique_times = entries.values_list(
            'start_time', 'end_time'
        ).distinct().order_by('start_time')
    else:
        unique_times = []
    
    # Create data for table
    data = []
    
    # Header row
    header = ['Time Slot'] + days
    data.append(header)
    
    # Add time slots and entries
    for start_time, end_time in unique_times:
        time_slot = f"{start_time.strftime('%I:%M %p')} - {end_time.strftime('%I:%M %p')}"
        row = [time_slot]
        
        for day in days:
            day_entries = entries.filter(
                day=day,
                start_time=start_time,
                end_time=end_time
            )
            
            cell_content = ""
            
            if day_entries.exists():
                # Check if it's a lab session (multiple entries with same subject)
                if day_entries.count() > 1:
                    subjects = list(day_entries.values_list('subject', flat=True).distinct())
                    if len(subjects) == 1:
                        # Lab session - show subject once, then all teachers
                        lab_subject = subjects[0]
                        lab_faculties = []
                        
                        for entry in day_entries:
                            if entry.faculty:
                                lab_faculties.append(entry.faculty.name)
                        
                        cell_content = f"<b>{lab_subject} (Lab)</b><br/>"
                        if lab_faculties:
                            for faculty in lab_faculties:
                                cell_content += f"{faculty}<br/>"
                        else:
                            cell_content += "No teachers assigned<br/>"
                    else:
                        # Different subjects (shouldn't happen)
                        for entry in day_entries:
                            subject_text = entry.subject
                            faculty_text = f"{entry.faculty.name}" if entry.faculty else "Break"
                            cell_content += f"<b>{subject_text}</b><br/>{faculty_text}<br/>"
                else:
                    # Single entry
                    entry = day_entries.first()
                    subject_text = entry.subject
                    faculty_text = f"{entry.faculty.name}" if entry.faculty else "Break"
                    cell_content = f"<b>{subject_text}</b><br/>{faculty_text}"
            
            if cell_content.strip():
                # Use Paragraph with HTML-like formatting
                cell_para = Paragraph(cell_content.strip(), 
                                     ParagraphStyle(
                                         'CellStyle',
                                         parent=styles['Normal'],
                                         fontSize=8,
                                         leading=9,
                                         spaceAfter=2
                                     ))
                row.append(cell_para)
            else:
                row.append("-")
        
        data.append(row)
    
    # Column widths
    col_widths = [120] + [115] * 6
    
    table = Table(data, colWidths=col_widths, repeatRows=1)
    
    # Style the table
    table_style = TableStyle([
        # Header style
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#475569')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('TOPPADDING', (0, 0), (-1, 0), 8),
        
        # Time column style
        ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#f8f9fa')),
        ('ALIGN', (0, 1), (0, -1), 'CENTER'),
        ('VALIGN', (0, 1), (0, -1), 'MIDDLE'),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 1), (0, -1), 9),
        
        # Grid lines
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
        ('VALIGN', (1, 1), (-1, -1), 'MIDDLE'),
        
        # Cell padding
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        
        # Row banding
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9fafb')]),
    ])
    
    table.setStyle(table_style)
    elements.append(table)
    elements.append(Spacer(1, 20))
    
    # Add signature section
    signature_style = ParagraphStyle(
        'SignatureStyle',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        spaceBefore=30
    )
    
    signature_data = [
        [Paragraph("HOD<br/><small>" + department.name + " Department</small>", signature_style),
         Paragraph("Director<br/><small>D.H.B. Soni College</small>", signature_style),
         Paragraph("Principal<br/><small>D.H.B. Soni College</small>", signature_style)],
        ["___________________", "___________________", "___________________"],
    ]
    
    signature_table = Table(signature_data, colWidths=[250, 250, 250])
    signature_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('FONTSIZE', (0, 1), (-1, 1), 12),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    
    elements.append(signature_table)
    
    # Footer
    footer_style = ParagraphStyle(
        'FooterStyle',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.gray,
        alignment=TA_CENTER,
        spaceBefore=20
    )
    
    elements.append(Paragraph(
        f"Generated on: {datetime.now().strftime('%d/%m/%Y at %I:%M %p')}",
        footer_style
    ))
    
    # Build PDF
    try:
        doc.build(elements)
    except Exception as e:
        return create_simple_pdf(buffer, department, selected_semester, entries)
    
    buffer.seek(0)
    filename = f"{department.name}_{selected_semester.replace(' ', '_')}_Timetable.pdf"
    
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def create_simple_pdf(buffer, department, selected_semester, entries):
    """Fallback PDF creation if main method fails"""
    from reportlab.lib.units import inch
    
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)
    
    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width/2, height-50, "D.H.B. SONI COLLEGE, SOLAPUR")
    
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(width/2, height-80, f"{department.name} Department")
    
    c.setFont("Helvetica", 12)
    c.drawCentredString(width/2, height-100, f"{selected_semester} Timetable")
    
    # Table header
    c.setFont("Helvetica-Bold", 10)
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    
    # Calculate column positions
    col_width = (width - 100) / 7
    x_positions = [40 + i * col_width for i in range(7)]
    
    # Draw headers
    headers = ['Time'] + days
    for i, header in enumerate(headers):
        c.drawString(x_positions[i] + 5, height-140, header[:10])
    
    # Draw table grid
    c.setStrokeColor(colors.black)
    
    # Horizontal lines
    for i in range(13):  # 12 rows max
        y = height - 150 - (i * 40)
        if y > 100:
            c.line(40, y, width-40, y)
    
    # Vertical lines
    for x in x_positions + [width-40]:
        c.line(x, height-150, x, max(100, height-150 - (12*40)))
    
    # Add data
    if entries.exists():
        unique_times = entries.values_list(
            'start_time', 'end_time'
        ).distinct().order_by('start_time')
        
        c.setFont("Helvetica", 8)
        row_height = 35
        current_y = height - 165
        
        for idx, (start_time, end_time) in enumerate(unique_times):
            if idx >= 12:  # Max 12 time slots
                break
                
            time_slot = f"{start_time.strftime('%I:%M')}-\n{end_time.strftime('%I:%M %p')}"
            
            # Draw time in first column
            c.drawString(x_positions[0] + 10, current_y, time_slot)
            
            # Draw entries for each day
            for day_idx, day in enumerate(days):
                day_entries = entries.filter(
                    day=day,
                    start_time=start_time,
                    end_time=end_time
                )
                
                text = ""
                for entry in day_entries:
                    subject = entry.subject[:15] + "..." if len(entry.subject) > 15 else entry.subject
                    faculty = entry.faculty.name[:10] if entry.faculty else "Break"
                    text += f"{subject}\n({faculty})\n"
                
                if text:
                    y = current_y
                    for line in text.split('\n'):
                        if line:
                            c.drawString(x_positions[day_idx+1] + 5, y, line)
                            y -= 10
            
            current_y -= row_height
    
    # Add signature section
    c.setFont("Helvetica", 10)
    c.drawString(80, 80, "___________________")
    c.drawString(80, 65, "HOD")
    c.drawString(80, 50, f"{department.name} Department")
    
    c.drawString(width/2 - 50, 80, "___________________")
    c.drawString(width/2 - 50, 65, "Director")
    c.drawString(width/2 - 50, 50, "D.H.B. Soni College")
    
    c.drawString(width - 180, 80, "___________________")
    c.drawString(width - 180, 65, "Principal")
    c.drawString(width - 180, 50, "D.H.B. Soni College")
    
    # Add generation date
    c.setFont("Helvetica", 8)
    c.drawString(width - 200, 30, f"Generated: {datetime.now().strftime('%d/%m/%Y %I:%M %p')}")
    
    c.showPage()
    c.save()
    
    return buffer
def download_timetable_excel(request, dept_id):
    """Generate Excel timetable with proper formatting"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    
    department = get_object_or_404(Department, id=dept_id)
    selected_semester = request.GET.get('semester', 'Semester 1')
    
    # Get entries
    entries = TimetableEntry.objects.filter(
        department=department,
        semester=selected_semester
    ).select_related('faculty')
    
    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = f"{selected_semester} Timetable"
    
    # Style definitions
    header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
    time_fill = PatternFill(start_color="E6E6E6", end_color="E6E6E6", fill_type="solid")
    header_font = Font(name='Calibri', size=12, bold=True, color="FFFFFF")
    time_font = Font(name='Calibri', size=11, bold=True)
    subject_font = Font(name='Calibri', size=10, bold=True)
    faculty_font = Font(name='Calibri', size=9)
    lab_font = Font(name='Calibri', size=10, bold=True, color="2E75B5")
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    cell_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    # Days
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    
    # Start from row 1
    current_row = 1
    
    # Title row
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=7)
    title_cell = ws.cell(row=current_row, column=1, value="D.H.B. SONI COLLEGE, SOLAPUR")
    title_cell.font = Font(name='Calibri', size=16, bold=True)
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    current_row += 1
    
    # Department row
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=7)
    dept_cell = ws.cell(row=current_row, column=1, value=f"{department.name} Department")
    dept_cell.font = Font(name='Calibri', size=14, bold=True, color="7030A0")
    dept_cell.alignment = Alignment(horizontal='center', vertical='center')
    current_row += 1
    
    # Semester row
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=7)
    semester_cell = ws.cell(row=current_row, column=1, value=f"{selected_semester} Timetable")
    semester_cell.font = Font(name='Calibri', size=12, bold=True, color="00B050")
    semester_cell.alignment = Alignment(horizontal='center', vertical='center')
    current_row += 1
    
    # Empty row
    current_row += 1
    
    # Table header
    headers = ['Time Slot'] + days
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=current_row, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_alignment
        cell.border = thin_border
    current_row += 1
    
    # Get unique time slots
    if entries.exists():
        unique_times = entries.values_list(
            'start_time', 'end_time'
        ).distinct().order_by('start_time')
    else:
        unique_times = []
    
    # Add data rows
    for start_time, end_time in unique_times:
        time_slot = f"{start_time.strftime('%I:%M %p')} - {end_time.strftime('%I:%M %p')}"
        
        # Time column cell
        time_cell = ws.cell(row=current_row, column=1, value=time_slot)
        time_cell.fill = time_fill
        time_cell.font = time_font
        time_cell.alignment = center_alignment
        time_cell.border = thin_border
        
        # Day columns
        for day_idx, day in enumerate(days, start=2):
            day_entries = entries.filter(
                day=day,
                start_time=start_time,
                end_time=end_time
            )
            
            if day_entries.exists():
                # Check if it's a lab session
                if day_entries.count() > 1:
                    subjects = list(day_entries.values_list('subject', flat=True).distinct())
                    if len(subjects) == 1:
                        # Lab session
                        lab_subject = subjects[0]
                        lab_faculties = []
                        
                        for entry in day_entries:
                            if entry.faculty:
                                lab_faculties.append(entry.faculty.name)
                        
                        # Create cell content with CHAR(10) for line breaks
                        cell_content = f"{lab_subject} (Lab)"
                        if lab_faculties:
                            for faculty in lab_faculties:
                                cell_content += f"\n{faculty}"
                        else:
                            cell_content += "\nNo teachers assigned"
                    else:
                        # Different subjects
                        cell_content = ""
                        for entry in day_entries:
                            subject_text = entry.subject
                            faculty_text = f"{entry.faculty.name}" if entry.faculty else "Break"
                            cell_content += f"{subject_text}\n{faculty_text}\n"
                else:
                    # Single entry
                    entry = day_entries.first()
                    subject_text = entry.subject
                    faculty_text = f"{entry.faculty.name}" if entry.faculty else "Break"
                    cell_content = f"{subject_text}\n{faculty_text}"
            else:
                cell_content = "-"
            
            cell = ws.cell(row=current_row, column=day_idx)
            cell.value = cell_content.strip()
            
            # Apply formatting
            if day_entries.exists() and day_entries.count() > 1:
                subjects = list(day_entries.values_list('subject', flat=True).distinct())
                if len(subjects) == 1:
                    cell.font = lab_font
                else:
                    cell.font = subject_font
            elif day_entries.exists():
                # Single entry - we can't format individual lines, so use subject font
                cell.font = subject_font
            
            cell.alignment = cell_alignment
            cell.border = thin_border
        
        current_row += 1
    
    # Set row heights for better visibility
    for row in range(5, current_row + 1):
        ws.row_dimensions[row].height = 40
    
    # Set column widths
    ws.column_dimensions['A'].width = 15  # Time column
    for col_idx in range(2, 8):  # Day columns B-G
        column_letter = get_column_letter(col_idx)
        ws.column_dimensions[column_letter].width = 20
    
    # Add signature section
    current_row += 2
    
    # Signature lines
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=2)
    ws.cell(row=current_row, column=1, value="___________________").alignment = Alignment(horizontal='center')
    
    ws.merge_cells(start_row=current_row, start_column=3, end_row=current_row, end_column=4)
    ws.cell(row=current_row, column=3, value="___________________").alignment = Alignment(horizontal='center')
    
    ws.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=6)
    ws.cell(row=current_row, column=5, value="___________________").alignment = Alignment(horizontal='center')
    
    current_row += 1
    
    # Signature titles
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=2)
    ws.cell(row=current_row, column=1, value="HOD").font = Font(bold=True)
    ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
    
    ws.merge_cells(start_row=current_row, start_column=3, end_row=current_row, end_column=4)
    ws.cell(row=current_row, column=3, value="Director").font = Font(bold=True)
    ws.cell(row=current_row, column=3).alignment = Alignment(horizontal='center')
    
    ws.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=6)
    ws.cell(row=current_row, column=5, value="Principal").font = Font(bold=True)
    ws.cell(row=current_row, column=5).alignment = Alignment(horizontal='center')
    
    current_row += 1
    
    # Signature details
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=2)
    ws.cell(row=current_row, column=1, value=f"{department.name} Department").font = Font(italic=True)
    ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
    
    ws.merge_cells(start_row=current_row, start_column=3, end_row=current_row, end_column=4)
    ws.cell(row=current_row, column=3, value="D.H.B. Soni College").font = Font(italic=True)
    ws.cell(row=current_row, column=3).alignment = Alignment(horizontal='center')
    
    ws.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=6)
    ws.cell(row=current_row, column=5, value="D.H.B. Soni College").font = Font(italic=True)
    ws.cell(row=current_row, column=5).alignment = Alignment(horizontal='center')
    
    # Add generation date
    current_row += 2
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=7)
    date_cell = ws.cell(row=current_row, column=1, 
                       value=f"Generated on: {datetime.now().strftime('%d/%m/%Y at %I:%M %p')}")
    date_cell.font = Font(italic=True, color="666666")
    date_cell.alignment = Alignment(horizontal='center')
    
    # Prepare response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    filename = f"{department.name}_{selected_semester.replace(' ', '_')}_Timetable.xlsx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    wb.save(response)
    return response
def download_timetable_csv(request, dept_id):
    """Generate CSV timetable with proper formatting"""
    department = get_object_or_404(Department, id=dept_id)
    selected_semester = request.GET.get('semester', 'Semester 1')
    
    # Get entries
    entries = TimetableEntry.objects.filter(
        department=department,
        semester=selected_semester
    ).select_related('faculty')
    
    # Create HttpResponse
    response = HttpResponse(content_type='text/csv')
    filename = f"{department.name}_{selected_semester.replace(' ', '_')}_Timetable.csv"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    # Create CSV writer
    writer = csv.writer(response)
    
    # Write headers
    writer.writerow(['D.H.B. SONI COLLEGE, SOLAPUR'])
    writer.writerow([f'{department.name} Department'])
    writer.writerow([f'{selected_semester} Timetable'])
    writer.writerow([])
    
    # Days
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    
    # Table header
    writer.writerow(['Time Slot'] + days)
    
    # Get unique time slots
    if entries.exists():
        unique_times = entries.values_list(
            'start_time', 'end_time'
        ).distinct().order_by('start_time')
    else:
        unique_times = []
    
    # Write data rows
    for start_time, end_time in unique_times:
        time_slot = f"{start_time.strftime('%I:%M %p')} - {end_time.strftime('%I:%M %p')}"
        row_data = [time_slot]
        
        for day in days:
            day_entries = entries.filter(
                day=day,
                start_time=start_time,
                end_time=end_time
            )
            
            cell_content = ""
            
            if day_entries.exists():
                # Check if it's a lab session
                if day_entries.count() > 1:
                    subjects = list(day_entries.values_list('subject', flat=True).distinct())
                    if len(subjects) == 1:
                        # Lab session
                        lab_subject = subjects[0]
                        lab_faculties = []
                        
                        for entry in day_entries:
                            if entry.faculty:
                                lab_faculties.append(entry.faculty.name)
                        
                        cell_content = f"{lab_subject} (Lab)"
                        if lab_faculties:
                            for faculty in lab_faculties:
                                cell_content += f"\\n{faculty}"
                        else:
                            cell_content += "\\nNo teachers assigned"
                    else:
                        # Different subjects
                        for entry in day_entries:
                            subject_text = entry.subject
                            faculty_text = f"{entry.faculty.name}" if entry.faculty else "Break"
                            cell_content += f"{subject_text}\\n{faculty_text}\\n"
                else:
                    # Single entry
                    entry = day_entries.first()
                    subject_text = entry.subject
                    faculty_text = f"{entry.faculty.name}" if entry.faculty else "Break"
                    cell_content = f"{subject_text}\\n{faculty_text}"
            else:
                cell_content = "-"
            
            row_data.append(cell_content.strip())
        
        writer.writerow(row_data)
    
    # Add empty rows
    writer.writerow([])
    writer.writerow([])
    
    # Add signature section
    writer.writerow(["", "", ""])
    writer.writerow(["___________________", "___________________", "___________________"])
    writer.writerow(["HOD", "Director", "Principal"])
    writer.writerow([f"{department.name} Department", "D.H.B. Soni College", "D.H.B. Soni College"])
    writer.writerow([])
    
    # Add generation date
    writer.writerow([f"Generated on: {datetime.now().strftime('%d/%m/%Y at %I:%M %p')}"])
    
    return response

def download_timetable_word(request, dept_id):
    """Generate Word document with proper formatting"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    
    department = get_object_or_404(Department, id=dept_id)
    selected_semester = request.GET.get('semester', 'Semester 1')
    
    # Get entries
    entries = TimetableEntry.objects.filter(
        department=department,
        semester=selected_semester
    ).select_related('faculty')
    
    # Create document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    
    # Title
    title = doc.add_paragraph('D.H.B. SONI COLLEGE, SOLAPUR')
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Department
    dept = doc.add_paragraph()
    dept_run = dept.add_run(f"{department.name} Department")
    dept_run.font.size = Pt(14)
    dept_run.font.bold = True
    dept_run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Semester
    sem = doc.add_paragraph()
    sem_run = sem.add_run(f"{selected_semester} Timetable")
    sem_run.font.size = Pt(12)
    sem_run.font.bold = True
    sem_run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    sem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Space
    doc.add_paragraph()
    
    # Create timetable table
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    
    if entries.exists():
        unique_times = entries.values_list(
            'start_time', 'end_time'
        ).distinct().order_by('start_time')
    else:
        unique_times = []
    
    # Create table
    num_rows = len(unique_times) + 1
    num_cols = len(days) + 1
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Cm(3.5)
    for i in range(1, num_cols):
        table.columns[i].width = Cm(4.0)
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Time Slot"
    header_cells[0].paragraphs[0].runs[0].font.bold = True
    header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, day in enumerate(days, 1):
        header_cells[i].text = day
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Fill data rows
    for row_idx, (start_time, end_time) in enumerate(unique_times, 1):
        time_slot = f"{start_time.strftime('%I:%M %p')} - {end_time.strftime('%I:%M %p')}"
        
        # Time cell
        time_cell = table.rows[row_idx].cells[0]
        time_cell.text = time_slot
        time_cell.paragraphs[0].runs[0].font.bold = True
        time_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Day cells
        for day_idx, day in enumerate(days, 1):
            cell = table.rows[row_idx].cells[day_idx]
            day_entries = entries.filter(
                day=day,
                start_time=start_time,
                end_time=end_time
            )
            
            # Clear any existing content
            cell.text = ""
            
            if day_entries.exists():
                # Check if it's a lab session
                if day_entries.count() > 1:
                    subjects = list(day_entries.values_list('subject', flat=True).distinct())
                    if len(subjects) == 1:
                        # Lab session
                        lab_subject = subjects[0]
                        lab_faculties = []
                        
                        for entry in day_entries:
                            if entry.faculty:
                                lab_faculties.append(entry.faculty.name)
                        
                        # Add subject line
                        subject_para = cell.add_paragraph()
                        subject_run = subject_para.add_run(f"{lab_subject} (Lab)")
                        subject_run.font.bold = True
                        subject_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Blue for lab
                        subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add faculty lines
                        if lab_faculties:
                            for faculty in lab_faculties:
                                faculty_para = cell.add_paragraph()
                                faculty_run = faculty_para.add_run(faculty)
                                faculty_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            faculty_para = cell.add_paragraph()
                            faculty_run = faculty_para.add_run("No teachers assigned")
                            faculty_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        # Different subjects
                        for entry in day_entries:
                            subject_text = entry.subject
                            faculty_text = f"{entry.faculty.name}" if entry.faculty else "Break"
                            
                            subject_para = cell.add_paragraph()
                            subject_run = subject_para.add_run(subject_text)
                            subject_run.font.bold = True
                            subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            faculty_para = cell.add_paragraph()
                            faculty_run = faculty_para.add_run(faculty_text)
                            faculty_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # Single entry
                    entry = day_entries.first()
                    subject_text = entry.subject
                    faculty_text = f"{entry.faculty.name}" if entry.faculty else "Break"
                    
                    subject_para = cell.add_paragraph()
                    subject_run = subject_para.add_run(subject_text)
                    subject_run.font.bold = True
                    subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    faculty_para = cell.add_paragraph()
                    faculty_run = faculty_para.add_run(faculty_text)
                    faculty_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Empty cell
                dash_para = cell.add_paragraph()
                dash_run = dash_para.add_run("-")
                dash_run.font.italic = True
                dash_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Space after table
    doc.add_paragraph()
    doc.add_paragraph()
    
    # SIGNATURE SECTION
    sig_table = doc.add_table(rows=3, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths for signatures
    for col in sig_table.columns:
        col.width = Cm(5.0)
    
    # Fill signature table
    signatures = [
        ("HOD", f"{department.name} Department"),
        ("Director", "D.H.B. Soni College"),
        ("Principal", "D.H.B. Soni College")
    ]
    
    for i, (title_text, detail) in enumerate(signatures):
        # Signature line (top row)
        sig_cell = sig_table.rows[0].cells[i]
        sig_cell.text = "___________________"
        sig_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title (middle row)
        title_cell = sig_table.rows[1].cells[i]
        title_cell.text = title_text
        title_cell.paragraphs[0].runs[0].font.bold = True
        title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Detail (bottom row)
        detail_cell = sig_table.rows[2].cells[i]
        detail_cell.text = detail
        detail_cell.paragraphs[0].runs[0].font.italic = True
        detail_cell.paragraphs[0].runs[0].font.size = Pt(9)
        detail_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Space before date
    doc.add_paragraph()
    
    # Generation date
    date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d/%m/%Y at %I:%M %p')}")
    date_para.runs[0].font.italic = True
    date_para.runs[0].font.size = Pt(9)
    date_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f"{department.name}_{selected_semester.replace(' ', '_')}_Timetable.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    doc.save(response)
    return response

def download_timetable_json(request, dept_id):
    """Generate JSON timetable for a specific department and semester"""
    department = get_object_or_404(Department, id=dept_id)
    selected_semester = request.GET.get('semester', 'Semester 1')
    
    # Get entries for this department and semester
    entries = TimetableEntry.objects.filter(
        department=department,
        semester=selected_semester
    ).select_related('faculty')
    
    # Prepare data structure
    data = {
        'college': 'D.H.B. SONI COLLEGE, SOLAPUR',
        'department': department.name,
        'semester': selected_semester,
        'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'timetable': {}
    }
    
    # Days of the week
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    
    # Get all unique time slots
    if entries.exists():
        unique_times = entries.values_list(
            'start_time', 'end_time'
        ).distinct().order_by('start_time')
    else:
        unique_times = []
    
    # Build timetable structure
    for start_time, end_time in unique_times:
        time_slot = f"{start_time.strftime('%H:%M')}-{end_time.strftime('%H:%M')}"
        data['timetable'][time_slot] = {}
        
        for day in days:
            day_entries = entries.filter(
                day=day,
                start_time=start_time,
                end_time=end_time
            )
            
            day_data = []
            for entry in day_entries:
                entry_data = {
                    'subject': entry.subject,
                    'faculty': entry.faculty.name if entry.faculty else None,
                    'faculty_id': entry.faculty.id if entry.faculty else None,
                    'is_break': not entry.faculty,
                    'entry_id': entry.id
                }
                day_data.append(entry_data)
            
            data['timetable'][time_slot][day] = day_data
    
    # Create HttpResponse with JSON
    response = HttpResponse(
        json.dumps(data, indent=2, default=str),
        content_type='application/json'
    )
    filename = f"{department.name}_{selected_semester.replace(' ', '_')}_Timetable.json"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response

def download_all_formats(request, dept_id):
    """Provide download options page"""
    department = get_object_or_404(Department, id=dept_id)
    selected_semester = request.GET.get('semester', 'Semester 1')
    
    has_entries = TimetableEntry.objects.filter(
        department=department,
        semester=selected_semester
    ).exists()
    
    return render(request, 'timetable/download_options.html', {
        'department': department,
        'selected_semester': selected_semester,
        'has_entries': has_entries,
    })

def share_timetable_image(request, dept_id):
    """Generate and share timetable as an image"""
    from PIL import Image, ImageDraw, ImageFont
    import io
    
    department = get_object_or_404(Department, id=dept_id)
    selected_semester = request.GET.get('semester', 'Semester 1')
    
    # Get entries
    entries = TimetableEntry.objects.filter(
        department=department,
        semester=selected_semester
    ).select_related('faculty')
    
    # Create image
    img_width = 1200
    img_height = 800
    img = Image.new('RGB', (img_width, img_height), color='white')
    draw = ImageDraw.Draw(img)
    
    # Try to load fonts
    try:
        title_font = ImageFont.truetype("arial.ttf", 36)
        header_font = ImageFont.truetype("arial.ttf", 24)
        text_font = ImageFont.truetype("arial.ttf", 16)
        small_font = ImageFont.truetype("arial.ttf", 12)
    except:
        # Fallback to default font
        title_font = ImageFont.load_default()
        header_font = ImageFont.load_default()
        text_font = ImageFont.load_default()
        small_font = ImageFont.load_default()
    
    # Colors
    title_color = (0, 0, 0)  # Black
    dept_color = (124, 58, 237)  # Purple
    sem_color = (16, 185, 129)  # Emerald
    header_bg = (71, 85, 105)  # Slate
    header_text = (255, 255, 255)  # White
    cell_bg = (248, 250, 252)  # Light gray
    text_color = (31, 41, 55)  # Dark gray
    border_color = (209, 213, 219)  # Gray
    
    # Helper function to get text width (compatible with older PIL)
    def get_text_width(text, font):
        try:
            # For newer PIL versions
            return draw.textlength(text, font=font)
        except AttributeError:
            # For older PIL versions - approximate width
            return len(text) * font.size // 2
    
    # Draw title
    title = "D.H.B. SONI COLLEGE, SOLAPUR"
    title_width = get_text_width(title, title_font)
    draw.text(((img_width - title_width) // 2, 30), title, fill=title_color, font=title_font)
    
    # Draw department
    dept_text = f"{department.name} Department"
    dept_width = get_text_width(dept_text, header_font)
    draw.text(((img_width - dept_width) // 2, 80), dept_text, fill=dept_color, font=header_font)
    
    # Draw semester
    sem_text = f"{selected_semester} Timetable"
    sem_width = get_text_width(sem_text, text_font)
    draw.text(((img_width - sem_width) // 2, 120), sem_text, fill=sem_color, font=text_font)
    
    # Days of the week
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    
    # Get unique time slots
    if entries.exists():
        unique_times = entries.values_list(
            'start_time', 'end_time'
        ).distinct().order_by('start_time')
    else:
        unique_times = []
    
    # Calculate table dimensions
    table_top = 180
    row_height = 60
    col_width = img_width // (len(days) + 1)
    
    # Draw table header
    for i, header in enumerate(['Time Slot'] + days):
        x1 = i * col_width
        y1 = table_top
        x2 = x1 + col_width
        y2 = y1 + row_height
        
        # Draw header cell
        draw.rectangle([x1, y1, x2, y2], fill=header_bg, outline=border_color, width=2)
        
        # Draw header text
        header_width = get_text_width(header, text_font)
        text_x = x1 + (col_width - header_width) // 2
        text_y = y1 + (row_height - 20) // 2
        draw.text((text_x, text_y), header, fill=header_text, font=text_font)
    
    # Draw data rows
    for row_idx, (start_time, end_time) in enumerate(unique_times):
        y1 = table_top + (row_idx + 1) * row_height
        y2 = y1 + row_height
        
        # Time cell
        x1 = 0
        x2 = col_width
        time_slot = f"{start_time.strftime('%I:%M %p')} - {end_time.strftime('%I:%M %p')}"
        
        # Draw time cell with light background
        draw.rectangle([x1, y1, x2, y2], fill=cell_bg, outline=border_color, width=1)
        
        # Draw time text
        time_width = get_text_width(time_slot, text_font)
        time_x = x1 + (col_width - time_width) // 2
        time_y = y1 + (row_height - 20) // 2
        draw.text((time_x, time_y), time_slot, fill=text_color, font=text_font)
        
        # Day cells
        for day_idx, day in enumerate(days, start=1):
            x1 = day_idx * col_width
            x2 = x1 + col_width
            
            # Draw cell border
            draw.rectangle([x1, y1, x2, y2], fill='white', outline=border_color, width=1)
            
            # Get entries for this cell
            day_entries = entries.filter(
                day=day,
                start_time=start_time,
                end_time=end_time
            )
            
            # Draw cell content
            if day_entries:
                for entry_idx, entry in enumerate(day_entries):
                    subject = entry.subject[:20] + "..." if len(entry.subject) > 20 else entry.subject
                    faculty = f"({entry.faculty.name[:15]})" if entry.faculty else "(Break)"
                    
                    # Draw subject
                    subject_width = get_text_width(subject, small_font)
                    subject_x = x1 + (col_width - subject_width) // 2
                    subject_y = y1 + 10 + (entry_idx * 40)
                    draw.text((subject_x, subject_y), subject, fill=text_color, font=small_font)
                    
                    # Draw faculty
                    faculty_width = get_text_width(faculty, small_font)
                    faculty_x = x1 + (col_width - faculty_width) // 2
                    faculty_y = subject_y + 20
                    draw.text((faculty_x, faculty_y), faculty, fill=text_color, font=small_font)
            else:
                # Draw dash for empty cell
                dash_width = get_text_width("-", text_font)
                dash_x = x1 + (col_width - dash_width) // 2
                dash_y = y1 + (row_height - 20) // 2
                draw.text((dash_x, dash_y), "-", fill=(156, 163, 175), font=text_font)
    
    # Draw college text at bottom
    bottom_text = f"{department.name} - {selected_semester} | D.H.B. Soni College, Solapur"
    bottom_width = get_text_width(bottom_text, small_font)
    draw.text(((img_width - bottom_width) // 2, img_height - 40), 
              bottom_text, fill=(107, 114, 128), font=small_font)
    
    # Add share URL text
    share_text = "Scan QR code or visit URL to view full timetable"
    share_width = get_text_width(share_text, small_font)
    draw.text(((img_width - share_width) // 2, img_height - 70), 
              share_text, fill=(59, 130, 246), font=small_font)
    
    # Save image to bytes
    img_buffer = io.BytesIO()
    img.save(img_buffer, format='PNG', quality=95)
    img_buffer.seek(0)
    
    # Create response
    response = HttpResponse(img_buffer, content_type='image/png')
    response['Content-Disposition'] = f'inline; filename="{department.name}_{selected_semester}_Timetable.png"'
    
    return response


def share_timetable_page(request, dept_id):
    """Page for sharing timetable with social media options"""
    department = get_object_or_404(Department, id=dept_id)
    selected_semester = request.GET.get('semester', 'Semester 1')
    
    # Build share URL
    base_url = request.build_absolute_uri('/').rstrip('/')
    timetable_url = f"{base_url}/timetable/{dept_id}/?semester={selected_semester}"
    image_url = f"{base_url}/timetable/{dept_id}/share-image/?semester={selected_semester}"
    
    # Social media share URLs
    whatsapp_url = f"https://wa.me/?text=Check%20out%20{department.name}%20{selected_semester}%20Timetable:%20{timetable_url}"
    telegram_url = f"https://t.me/share/url?url={timetable_url}&text={department.name}%20{selected_semester}%20Timetable"
    facebook_url = f"https://www.facebook.com/sharer/sharer.php?u={timetable_url}"
    twitter_url = f"https://twitter.com/intent/tweet?url={timetable_url}&text={department.name}%20{selected_semester}%20Timetable"
    
    return render(request, 'timetable/share_timetable.html', {
        'department': department,
        'selected_semester': selected_semester,
        'timetable_url': timetable_url,
        'image_url': image_url,
        'whatsapp_url': whatsapp_url,
        'telegram_url': telegram_url,
        'facebook_url': facebook_url,
        'twitter_url': twitter_url,
    })

# timetable/views.py

# timetable/views.py

def check_faculty_conflict(faculty, day, start_time, end_time, current_dept, exclude_id=None, is_lab=False):
    """
    Checks if a faculty is busy in ANY department's currently active semester.
    """
    # 1. Find all potential time overlaps across the whole college
    potential_overlaps = TimetableEntry.objects.filter(
        faculty=faculty,
        day=day,
        start_time__lt=end_time,
        end_time__gt=start_time
    )

    if exclude_id:
        potential_overlaps = potential_overlaps.exclude(id=exclude_id)

    active_conflicts = []
    for entry in potential_overlaps:
        # 2. KEY LOGIC: Only count it as a conflict if the entry belongs 
        # to the semester currently marked 'Active' for THAT department.
        if entry.semester == entry.department.active_semester:
            
            # 3. Lab Exception: Allow multiple teachers in the SAME department
            if is_lab and entry.department == current_dept:
                continue
                
            active_conflicts.append(entry)
    
    return active_conflicts
# timetable/views.py

#@login_required
def timetable_create(request):
    if request.method == 'POST':
        form = TimetableForm(request.POST)
        
        # Check the radio button value from the template
        is_lab_mode = request.POST.get('is_lab') == 'lab'
        
        if form.is_valid():
            try:
                department = form.cleaned_data['department']
                day = form.cleaned_data['day']
                start_time = form.cleaned_data['start_time']
                end_time = form.cleaned_data['end_time']
                subject = form.cleaned_data['subject']
                semester = form.cleaned_data['semester']
                
                subject_lower = subject.lower()
                is_break = any(x in subject_lower for x in ['recess', 'lunch', 'break'])
                
                # --- CHECK FOR DUPLICATE TIMESLOT IN SAME DEPT/SEM ---
                existing_entries = TimetableEntry.objects.filter(
                    department=department,
                    semester=semester,
                    day=day,
                    start_time=start_time,
                    end_time=end_time
                )
                
                # If not lab mode, prevent multiple entries in the same slot
                if not is_lab_mode and existing_entries.exists():
                    messages.error(request, 
                        f'This time slot is already occupied in {department.name} - {semester}. '
                        'For multiple teachers, please select "Lab Session".'
                    )
                    return render(request, 'timetable/timetable_create.html', {'form': form})
                
                # --- LOGIC FOR LAB MODE (Multiple Teachers) ---
                if is_lab_mode:
                    lab_faculty = form.cleaned_data.get('lab_faculty')
                    
                    if not lab_faculty:
                        messages.error(request, "Please select at least one teacher for the Lab session.")
                        return render(request, 'timetable/timetable_create.html', {'form': form})

                    for faculty_member in lab_faculty:
                        # Call global conflict checker
                        conflicts = check_faculty_conflict(
                            faculty_member, day, start_time, end_time, 
                            current_dept=department, is_lab=True
                        )
                        if conflicts:
                            conflict = conflicts[0]
                            messages.error(request, 
                                f'CONFLICT: {faculty_member.name} is already assigned to '
                                f'{conflict.department.name} ({conflict.semester})'
                            )
                            return render(request, 'timetable/timetable_create.html', {'form': form})
                    
                    # Create entries for each lab teacher
                    for faculty_member in lab_faculty:
                        TimetableEntry.objects.create(
                            department=department, semester=semester, day=day,
                            start_time=start_time, end_time=end_time,
                            subject=subject, faculty=faculty_member
                        )
                    messages.success(request, f'Lab session created successfully!')

                # --- LOGIC FOR LECTURE MODE (Single Teacher) ---
                else:
                    faculty = form.cleaned_data.get('faculty')
                    
                    if is_break:
                        # Break Logic: Create for specific day or all days if day is empty
                        if day:
                            TimetableEntry.objects.create(
                                department=department, semester=semester, day=day,
                                start_time=start_time, end_time=end_time,
                                subject=subject, faculty=None
                            )
                        else:
                            days_list = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
                            for d in days_list:
                                TimetableEntry.objects.get_or_create(
                                    department=department, semester=semester, day=d,
                                    start_time=start_time, end_time=end_time,
                                    defaults={'subject': subject, 'faculty': None}
                                )
                        messages.success(request, f'Break/Recess added!')
                    else:
                        if faculty:
                            # Global conflict check for single teacher
                            conflicts = check_faculty_conflict(
                                faculty, day, start_time, end_time, 
                                current_dept=department, is_lab=False
                            )
                            if conflicts:
                                conflict = conflicts[0]
                                messages.error(request, 
                                    f'CONFLICT: {faculty.name} is in {conflict.department.name} ({conflict.semester})'
                                )
                                return render(request, 'timetable/timetable_create.html', {'form': form})
                        
                        # Save standard entry
                        entry = form.save(commit=False)
                        entry.save()
                        messages.success(request, f'Entry added for {semester}!')

                return redirect(f'/timetable/{department.id}/?semester={semester}')
                
            except Exception as e:
                messages.error(request, f'Error: {str(e)}')
                return render(request, 'timetable/timetable_create.html', {'form': form})
        else:
            messages.error(request, 'Please correct the validation errors.')
            
    else:
        # GET Request logic
        form = TimetableForm()
        dept_id = request.GET.get('department')
        semester = request.GET.get('semester')
        if dept_id:
            form.fields['department'].initial = Department.objects.filter(id=dept_id).first()
        if semester:
            form.fields['semester'].initial = semester
    
    return render(request, 'timetable/timetable_create.html', {'form': form})

def delete_entry(request, entry_id):
    entry = get_object_or_404(TimetableEntry, id=entry_id)
    department_id = entry.department.id
    semester = entry.semester
    entry.delete()
    messages.success(request, 'Entry deleted!')
    
    # Redirect back to same semester view
    return redirect(f'/timetable/{department_id}/?semester={semester}')
def department_history(request, dept_id):
    """View to list all archived years for a department"""
    department = get_object_or_404(Department, id=dept_id)
    
    # Order by year descending, then by created_at descending
    history_records = TimetableHistory.objects.filter(
        department=department
    ).order_by('-year', '-created_at')
    
    # For search bar
    all_departments = Department.objects.all()
    all_years = TimetableHistory.objects.values('year').distinct().order_by('-year')
    
    return render(request, 'timetable/history_list.html', {
        'department': department,
        'history_records': history_records,
        'all_departments': all_departments,
        'all_years': all_years,
    })

def archive_current_timetable(request, dept_id):
    department = get_object_or_404(Department, id=dept_id)
    semester = request.GET.get('semester', 'Semester 1')
    current_year = 2026  # Your target year
    
    # 1. Get all entries currently in this timetable
    entries = TimetableEntry.objects.filter(department=department, semester=semester)
    
    if not entries.exists():
        messages.warning(request, f"No entries found to archive for {semester}.")
        return redirect(f'/timetable/{dept_id}/?semester={semester}')

    # 2. Prepare the data snapshot
    data_list = []
    for entry in entries:
        data_list.append({
            'subject': entry.subject,
            'faculty': entry.faculty.name if entry.faculty else "Break/Recess",
            'day': entry.day,
            'start_time': entry.start_time.strftime('%I:%M %p'),
            'end_time': entry.end_time.strftime('%I:%M %p'),
        })

    # 3. SIMPLE FIX: ALWAYS CREATE NEW RECORD WITHOUT VERSION CHECK
    TimetableHistory.objects.create(
        department=department,
        semester=semester,
        year=current_year,
        data_snapshot=data_list
    )

    messages.success(request, f"Timetable for {semester} ({current_year}) has been archived successfully!")
    return redirect('department_history', dept_id=dept_id)
# In your views.py (wherever you have history views)
def view_history_detail(request, record_id):
    """View to display a specific archived timetable from history"""
    record = get_object_or_404(TimetableHistory, id=record_id)
    department = record.department
    selected_semester = record.semester
    
    # Days of the week
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    
    # Get data from history
    matrix_data = record.data_snapshot
    
    # Organize data into matrix format with lab session grouping
    time_slots_dict = {}
    
    for entry in matrix_data:
        time_key = f"{entry['start_time']} - {entry['end_time']}"
        if time_key not in time_slots_dict:
            time_slots_dict[time_key] = {day: [] for day in days}
        
        # Add entry to the correct day
        if entry['day'] in time_slots_dict[time_key]:
            time_slots_dict[time_key][entry['day']].append(entry)
    
    # Convert to matrix format for template with lab session grouping
    matrix = []
    for time_slot, day_data in sorted(time_slots_dict.items()):
        row_data = []
        for day in days:
            day_entries = day_data[day]
            
            # Group lab sessions
            if len(day_entries) > 1:
                # Check if these are lab sessions (same subject, multiple teachers)
                subjects = list(set([entry['subject'] for entry in day_entries]))
                if len(subjects) == 1:
                    # It's a lab session with multiple teachers
                    lab_subject = subjects[0]
                    lab_faculties = [entry['faculty'] for entry in day_entries if entry['faculty'] != 'Break/Recess']
                    
                    # Create a lab entry dictionary
                    lab_entry = {
                        'is_lab': True,
                        'subject': lab_subject,
                        'faculties': lab_faculties,
                        'faculty': None,
                    }
                    row_data.append([lab_entry])  # Put in a list to match template
                else:
                    # Different subjects at same time
                    row_data.append(day_entries)
            else:
                row_data.append(day_entries)
        
        matrix.append({
            'time': time_slot,
            'data': row_data
        })
    
    # ADD THESE LINES FOR SEARCH BAR
    all_departments = Department.objects.all()
    all_years = TimetableHistory.objects.values('year').distinct().order_by('-year')
    
    return render(request, 'timetable/history_detail.html', {
        'record': record,
        'department': department,
        'selected_semester': selected_semester,
        'days': days,
        'matrix': matrix,
        'has_entries': len(matrix_data) > 0,
        # ADD THESE FOR SEARCH BAR
        'all_departments': all_departments,
        'all_years': all_years,
    })
def archive_search(request):
    """Show search results on the same search page"""
    if request.method == 'GET':
        department_id = request.GET.get('department_id')
        year = request.GET.get('year')
        semester = request.GET.get('semester')
        
        # Get all departments and years for the search form
        all_departments = Department.objects.all()
        all_years = TimetableHistory.objects.values('year').distinct().order_by('-year')
        
        # If no search criteria, just show empty search page
        if not all([department_id, year, semester]):
            return render(request, 'timetable/archive_search_results.html', {
                'all_departments': all_departments,
                'all_years': all_years,
                'search_results': [],
            })
        
        # Find matching records
        department = get_object_or_404(Department, id=department_id)
        records = TimetableHistory.objects.filter(
            department_id=department_id,
            year=year,
            semester__icontains=semester
        ).order_by('-created_at')
        
        # Show search results on the same page
        return render(request, 'timetable/archive_search_results.html', {
            'department': department,
            'search_year': year,
            'search_semester': semester,
            'search_results': records,
            'all_departments': all_departments,
            'all_years': all_years,
        })
    
    return redirect('department_list')

def delete_all_entries(request, dept_id):
    """Delete all timetable entries for a specific department and semester"""
    department = get_object_or_404(Department, id=dept_id)
    selected_semester = request.GET.get('semester', 'Semester 1')
    
    if request.method == 'POST':
        # Delete all entries for this department and semester
        deleted_count, _ = TimetableEntry.objects.filter(
            department=department,
            semester=selected_semester
        ).delete()
        
        messages.success(request, f'Deleted {deleted_count} entries from {selected_semester}!')
        return redirect(f'/timetable/{dept_id}/?semester={selected_semester}')
    
    # If GET request, show confirmation page
    entry_count = TimetableEntry.objects.filter(
        department=department,
        semester=selected_semester
    ).count()
    
    return render(request, 'timetable/delete_all_confirm.html', {
        'department': department,
        'selected_semester': selected_semester,
        'entry_count': entry_count,
    })

def download_timetable_image(request, dept_id):
    """Download timetable as an image file with proper tabular format and enhanced lab session display"""
    department = get_object_or_404(Department, id=dept_id)
    selected_semester = request.GET.get('semester', 'Semester 1')
    
    # Get entries
    entries = TimetableEntry.objects.filter(
        department=department,
        semester=selected_semester
    ).select_related('faculty')
    
    # Create image dimensions - increased height for better lab session display
    img_width = 1200
    img_height = 1200  # Increased to accommodate lab session data
    
    from PIL import Image, ImageDraw, ImageFont
    import io
    
    img = Image.new('RGB', (img_width, img_height), color='white')
    draw = ImageDraw.Draw(img)
    
    # Fonts
    try:
        title_font = ImageFont.truetype("arial.ttf", 32)
        dept_font = ImageFont.truetype("arial.ttf", 26)
        sem_font = ImageFont.truetype("arial.ttf", 22)
        header_font = ImageFont.truetype("arial.ttf", 18)
        cell_font = ImageFont.truetype("arial.ttf", 14)
        small_font = ImageFont.truetype("arial.ttf", 11)
        signature_font = ImageFont.truetype("arial.ttf", 16)
        signature_small_font = ImageFont.truetype("arial.ttf", 12)
    except:
        title_font = ImageFont.load_default()
        dept_font = ImageFont.load_default()
        sem_font = ImageFont.load_default()
        header_font = ImageFont.load_default()
        cell_font = ImageFont.load_default()
        small_font = ImageFont.load_default()
        signature_font = ImageFont.load_default()
        signature_small_font = ImageFont.load_default()
    
    # Colors
    title_color = (0, 0, 0)
    dept_color = (124, 58, 237)  # Purple
    sem_color = (16, 185, 129)   # Emerald
    header_bg = (71, 85, 105)    # Slate-600 (matching web view)
    header_text = (255, 255, 255)
    time_bg = (248, 250, 252)    # Slate-50
    cell_bg = (255, 255, 255)
    lab_bg = (239, 246, 255)     # Light blue for lab sessions
    lab_color = (0, 51, 153)     # DARKER Blue for lab text
    border_color = (0, 0, 0)     # BLACK for clear borders
    text_color = (0, 0, 0)       # Black for text
    faculty_color = (0, 0, 0)    # BLACK for faculty names
    break_color = (100, 100, 100)  # Dark gray for breaks
    grid_color = (0, 0, 0)       # BLACK for all grid lines
    
    # Line widths - INCREASED FOR BETTER VISIBILITY
    outer_border_width = 5
    header_border_width = 4
    grid_line_width = 3  # Increased from 2 to 3
    
    def get_text_width(text, font):
        try:
            return draw.textlength(text, font=font)
        except AttributeError:
            return len(text) * font.size // 2
    
    # Title - CENTERED
    title = "D.H.B. SONI COLLEGE, SOLAPUR"
    title_width = get_text_width(title, title_font)
    draw.text(((img_width - title_width) // 2, 30), title, fill=title_color, font=title_font)
    
    # Department - CENTERED
    dept_text = f"{department.name} Department"
    dept_width = get_text_width(dept_text, dept_font)
    draw.text(((img_width - dept_width) // 2, 80), dept_text, fill=dept_color, font=dept_font)
    
    # Semester - CENTERED
    sem_text = f"{selected_semester} Timetable"
    sem_width = get_text_width(sem_text, sem_font)
    draw.text(((img_width - sem_width) // 2, 120), sem_text, fill=sem_color, font=sem_font)
    
    # Days
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    
    # Get time slots
    if entries.exists():
        unique_times = entries.values_list('start_time', 'end_time').distinct().order_by('start_time')
        time_slots = list(unique_times)
    else:
        time_slots = []
    
    # Table position
    table_top = 180
    
    # Table dimensions - increased row height for lab sessions
    row_height = 100
    time_col_width = 200
    day_col_width = (img_width - time_col_width - 100) // 6
    
    # Calculate actual table bottom based on number of time slots
    header_height = 60
    table_bottom = table_top + header_height + (len(time_slots) * row_height)
    
    # Draw outer table border (THICK BLACK BORDER)
    draw.rectangle(
        [50-outer_border_width, table_top-outer_border_width, 
         img_width-50+outer_border_width, table_bottom+outer_border_width],
        outline=border_color,
        width=outer_border_width
    )
    
    # Header row
    y_header = table_top
    
    # Header background
    draw.rectangle([50, y_header, img_width-50, y_header+header_height], fill=header_bg)
    
    # Time header cell
    time_header = "Time Slot"
    time_w = get_text_width(time_header, header_font)
    time_x = 50 + (time_col_width // 2) - (time_w // 2)
    time_y = y_header + (header_height // 2) - 10
    draw.text((time_x, time_y), time_header, fill=header_text, font=header_font)
    
    # Draw THICK vertical line after time column
    draw.line([50 + time_col_width, y_header, 50 + time_col_width, table_bottom], 
             fill=grid_color, width=grid_line_width)
    
    # Day headers with vertical lines
    x_day = 50 + time_col_width
    for i, day in enumerate(days):
        day_w = get_text_width(day, header_font)
        day_x = x_day + (day_col_width // 2) - (day_w // 2)
        day_y = y_header + (header_height // 2) - 10
        draw.text((day_x, day_y), day, fill=header_text, font=header_font)
        
        x_day += day_col_width
    
    # Draw THICK horizontal line after header
    draw.line([50, y_header+header_height, img_width-50, y_header+header_height], 
             fill=grid_color, width=header_border_width)
    
    # Data rows
    current_y = y_header + header_height
    
    for row_idx, (start_time, end_time) in enumerate(time_slots):
        
        # Time cell background
        time_cell_bg = time_bg
        draw.rectangle([50, current_y, 50+time_col_width, current_y+row_height], fill=time_cell_bg)
        
        # Time text
        time_text = f"{start_time.strftime('%I:%M %p')} - {end_time.strftime('%I:%M %p')}"
        time_w = get_text_width(time_text, cell_font)
        time_x = 50 + (time_col_width // 2) - (time_w // 2)
        time_y = current_y + (row_height // 2) - 10
        draw.text((time_x, time_y), time_text, fill=text_color, font=cell_font)
        
        # Day cells
        x_cell = 50 + time_col_width
        
        for day_idx, day in enumerate(days):
            cell_x1 = x_cell
            cell_x2 = cell_x1 + day_col_width
            
            # Get entries for this cell
            day_entries = entries.filter(day=day, start_time=start_time, end_time=end_time)
            
            # Check for lab session
            is_lab = False
            if day_entries.count() > 1:
                subjects = list(set([e.subject for e in day_entries]))
                if len(subjects) == 1:
                    is_lab = True
            
            # Cell background
            if is_lab:
                cell_bg_color = lab_bg
            else:
                cell_bg_color = cell_bg
            
            draw.rectangle([cell_x1, current_y, cell_x2, current_y+row_height], fill=cell_bg_color)
            
            # Draw cell content
            if day_entries.exists():
                y_offset = current_y + 8
                
                if is_lab:
                    # LAB SESSION - Show all details
                    lab_subject = day_entries.first().subject
                    lab_faculties = [e.faculty.name for e in day_entries if e.faculty]
                    
                    # LAB Label
                    lab_label = "LAB SESSION"
                    lab_label_w = get_text_width(lab_label, small_font)
                    lab_label_x = cell_x1 + (day_col_width // 2) - (lab_label_w // 2)
                    draw.text((lab_label_x, y_offset), lab_label, fill=lab_color, font=small_font)
                    y_offset += 15
                    
                    # Subject
                    subject_text = f"{lab_subject}"
                    if len(subject_text) > 20:
                        subject_text = subject_text[:18] + "..."
                    subj_w = get_text_width(subject_text, cell_font)
                    subj_x = cell_x1 + (day_col_width // 2) - (subj_w // 2)
                    draw.text((subj_x, y_offset), subject_text, fill=text_color, font=cell_font)
                    y_offset += 20
                    
                    # Faculty List
                    if lab_faculties:
                        for i, faculty in enumerate(lab_faculties[:3]):
                            faculty_text = f"({faculty})"
                            if len(faculty_text) > 25:
                                faculty_text = faculty_text[:23] + "..."
                            fac_w = get_text_width(faculty_text, small_font)
                            fac_x = cell_x1 + (day_col_width // 2) - (fac_w // 2)
                            draw.text((fac_x, y_offset), faculty_text, 
                                     fill=faculty_color, font=small_font)
                            y_offset += 13
                else:
                    # Single entry
                    entry = day_entries.first()
                    
                    # Subject
                    subject = entry.subject
                    if len(subject) > 20:
                        subject = subject[:18] + "..."
                    subj_w = get_text_width(subject, cell_font)
                    subj_x = cell_x1 + (day_col_width // 2) - (subj_w // 2)
                    draw.text((subj_x, y_offset), subject, fill=text_color, font=cell_font)
                    y_offset += 25
                    
                    # Faculty or Break
                    if entry.faculty:
                        faculty = f"({entry.faculty.name})"
                        if len(faculty) > 20:
                            faculty = faculty[:18] + "..."
                        fac_w = get_text_width(faculty, small_font)
                        fac_x = cell_x1 + (day_col_width // 2) - (fac_w // 2)
                        draw.text((fac_x, y_offset), faculty, fill=faculty_color, font=small_font)
                    else:
                        fac_w = get_text_width("(Break)", small_font)
                        fac_x = cell_x1 + (day_col_width // 2) - (fac_w // 2)
                        draw.text((fac_x, y_offset), "(Break)", fill=break_color, font=small_font)
            else:
                # Empty cell
                dash = "-"
                dash_w = get_text_width(dash, cell_font)
                dash_x = cell_x1 + (day_col_width // 2) - (dash_w // 2)
                dash_y = current_y + (row_height // 2) - 10
                draw.text((dash_x, dash_y), dash, fill=(156, 163, 175), font=cell_font)
            
            x_cell += day_col_width
        
        current_y += row_height
    
    # NOW DRAW ALL GRID LINES (after cells so they appear on top)
    
    # Draw ALL VERTICAL LINES from header to bottom
    # Line after time column
    draw.line([50 + time_col_width, table_top, 50 + time_col_width, table_bottom], 
             fill=grid_color, width=grid_line_width)
    
    # Lines between each day column
    for i in range(len(days)):
        x_pos = 50 + time_col_width + ((i + 1) * day_col_width)
        draw.line([x_pos, table_top, x_pos, table_bottom], 
                 fill=grid_color, width=grid_line_width)
    
    # Draw horizontal line after each time slot row
    y_line = y_header + header_height
    for idx in range(len(time_slots) + 1):  # +1 to include final bottom line
        draw.line([50, y_line, img_width-50, y_line], 
                 fill=grid_color, width=grid_line_width)
        y_line += row_height
    
    # SIGNATURE SECTION
    sig_y = table_bottom + 60
    sig_line_length = 200
    sig_spacing = (img_width - 300) // 3
    
    # HOD
    hod_x = 150
    draw.line([hod_x, sig_y, hod_x + sig_line_length, sig_y], fill=(100, 100, 100), width=2)
    hod_text = "Head of Department"
    hod_w = get_text_width(hod_text, signature_font)
    draw.text((hod_x + (sig_line_length//2) - (hod_w//2), sig_y + 15), hod_text, 
             fill=(80, 80, 80), font=signature_font)
    hod_dept = f"{department.name} Department"
    hod_dept_w = get_text_width(hod_dept, signature_small_font)
    draw.text((hod_x + (sig_line_length//2) - (hod_dept_w//2), sig_y + 40), hod_dept, 
             fill=(120, 120, 120), font=signature_small_font)
    
    # Director
    dir_x = hod_x + sig_spacing
    draw.line([dir_x, sig_y, dir_x + sig_line_length, sig_y], fill=(100, 100, 100), width=2)
    dir_text = "Director"
    dir_w = get_text_width(dir_text, signature_font)
    draw.text((dir_x + (sig_line_length//2) - (dir_w//2), sig_y + 15), dir_text, 
             fill=(80, 80, 80), font=signature_font)
    dir_dept = "D.H.B. Soni College"
    dir_dept_w = get_text_width(dir_dept, signature_small_font)
    draw.text((dir_x + (sig_line_length//2) - (dir_dept_w//2), sig_y + 40), dir_dept, 
             fill=(120, 120, 120), font=signature_small_font)
    
    # Principal
    prin_x = dir_x + sig_spacing
    draw.line([prin_x, sig_y, prin_x + sig_line_length, sig_y], fill=(100, 100, 100), width=2)
    prin_text = "Principal"
    prin_w = get_text_width(prin_text, signature_font)
    draw.text((prin_x + (sig_line_length//2) - (prin_w//2), sig_y + 15), prin_text, 
             fill=(80, 80, 80), font=signature_font)
    prin_dept = "D.H.B. Soni College"
    prin_dept_w = get_text_width(prin_dept, signature_small_font)
    draw.text((prin_x + (sig_line_length//2) - (prin_dept_w//2), sig_y + 40), prin_dept, 
             fill=(120, 120, 120), font=signature_small_font)
    
    # Footer
    footer_y = sig_y + 90
    footer_text = f"Generated on: {datetime.now().strftime('%d/%m/%Y at %I:%M %p')}"
    footer_w = get_text_width(footer_text, small_font)
    draw.text(((img_width - footer_w) // 2, footer_y), footer_text, 
             fill=(150, 150, 150), font=small_font)
    
    # College name at bottom
    bottom_y = img_height - 40
    bottom_text = f"BCS-II - {selected_semester} | D.H.B. Soni College, Solapur"
    bottom_w = get_text_width(bottom_text, small_font)
    draw.text(((img_width - bottom_w) // 2, bottom_y), bottom_text, 
             fill=(180, 180, 180), font=small_font)
    
    # Save image
    img_buffer = io.BytesIO()
    img.save(img_buffer, format='PNG', quality=100, optimize=True)
    img_buffer.seek(0)
    
    # Response
    filename = f"{department.name}_{selected_semester.replace(' ', '_')}_Timetable.png"
    response = HttpResponse(img_buffer, content_type='image/png')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response   
def set_semester_active(request, dept_id):
    department = get_object_or_404(Department, id=dept_id)
    selected_semester = request.GET.get('semester')

    if selected_semester:
        # Update the Department model's active_semester field
        department.active_semester = selected_semester
        department.save()
        
        messages.success(request, f"{selected_semester} is now the ACTIVE timetable for {department.name}.")
    
    # Redirect back to the timetable view for that semester
    return redirect(f'/timetable/{dept_id}/?semester={selected_semester}')

 